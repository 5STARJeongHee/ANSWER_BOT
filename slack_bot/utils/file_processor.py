# 다양한 형식의 Slack 첨부파일에서 텍스트를 추출하는 유틸리티
# 지원 형식: txt/trf/log/csv 등 텍스트 계열, xlsx/xls, docx/doc, pdf, mov/mp4 등 동영상
from __future__ import annotations
import io
import logging
from typing import Optional

import requests

from utils.attachment_result import AttachmentResult

logger = logging.getLogger(__name__)

_DOWNLOAD_TIMEOUT_SEC = 30
_MAX_TEXT_CHARS = 3000     # 파일 1개당 최대 추출 문자 수 (RAG 토큰 과부하 방지)
_MAX_XLSX_ROWS = 200       # xlsx 최대 읽기 행 수
_MAX_XLSX_COLS = 20        # xlsx 최대 읽기 열 수

# ---- MIME 타입 분류 ----
_TEXT_MIME_PREFIXES = ("text/",)

_XLSX_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}

_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/vnd.ms-word",
}

_PDF_MIMES = {"application/pdf"}

_VIDEO_MIMES = {
    "video/quicktime",   # .mov
    "video/mp4",
    "video/mpeg",
    "video/x-msvideo",  # .avi
    "video/webm",
}

# 확장자 기반 fallback (MIME=application/octet-stream 등 미지정 파일 대응)
_TEXT_EXTENSIONS = {".txt", ".trf", ".log", ".csv", ".tsv", ".md", ".yaml", ".yml", ".json", ".xml", ".ini", ".cfg"}
_XLSX_EXTENSIONS = {".xlsx", ".xls"}
_DOCX_EXTENSIONS = {".docx", ".doc"}
_PDF_EXTENSIONS  = {".pdf"}
_VIDEO_EXTENSIONS = {".mov", ".mp4", ".mpeg", ".mpg", ".avi", ".webm", ".mkv"}


def _get_extension(filename: str) -> str:
    """파일명에서 소문자 확장자를 반환한다."""
    if "." in filename:
        return "." + filename.rsplit(".", 1)[-1].lower()
    return ""


def _download_bytes(url: str, bot_token: str) -> Optional[bytes]:
    """Slack 비공개 URL에서 파일 바이트를 다운로드한다."""
    try:
        resp = requests.get(
            url,
            headers={"Authorization": f"Bearer {bot_token}"},
            timeout=_DOWNLOAD_TIMEOUT_SEC,
        )
        resp.raise_for_status()
        return resp.content
    except Exception as exc:
        logger.warning(f"파일 다운로드 실패 ({url[:60]}...): {exc}")
        return None


def _extract_text_plain(data: bytes) -> str:
    """텍스트 파일(txt, trf, log, csv 등)에서 텍스트를 추출한다."""
    for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_xlsx(data: bytes) -> str:
    """xlsx 파일의 시트 내용을 탭 구분 텍스트로 추출한다."""
    try:
        import openpyxl
    except ImportError:
        logger.warning("openpyxl 미설치 — xlsx 추출 불가. pip install openpyxl")
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_lines = [f"[시트: {sheet_name}]"]
            for row in ws.iter_rows(max_row=_MAX_XLSX_ROWS, max_col=_MAX_XLSX_COLS, values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    sheet_lines.append("\t".join(cells))
            if len(sheet_lines) > 1:
                parts.append("\n".join(sheet_lines))
        wb.close()
        return "\n\n".join(parts)
    except Exception as exc:
        logger.warning(f"xlsx 추출 실패: {exc}")
        return ""


def _extract_docx(data: bytes) -> str:
    """docx/doc 파일에서 단락 및 표 텍스트를 추출한다."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx 미설치 — docx 추출 불가. pip install python-docx")
        return ""
    try:
        doc = Document(io.BytesIO(data))
        lines: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as exc:
        logger.warning(f"docx 추출 실패: {exc}")
        return ""


def _extract_pdf(data: bytes) -> str:
    """pdf 파일에서 텍스트를 추출한다."""
    try:
        import pypdf
    except ImportError:
        logger.warning("pypdf 미설치 — pdf 추출 불가. pip install pypdf")
        return ""
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() for page in reader.pages]
        return "\n".join(t.strip() for t in pages if t)
    except Exception as exc:
        logger.warning(f"pdf 추출 실패: {exc}")
        return ""


def _extract_video(data: bytes, filename: str) -> str:
    """
    동영상 파일을 openai-whisper로 전사한다 (선택 설치).
    whisper 미설치 시 파일명만 기록하고 반환한다.
    """
    try:
        import whisper
        import tempfile
        import os
    except ImportError:
        logger.info(f"openai-whisper 미설치 — 동영상 전사 생략 ({filename}). pip install openai-whisper")
        return f"[동영상 첨부: {filename}]"
    try:
        ext = _get_extension(filename) or ".mp4"
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            model = whisper.load_model("base")
            result = model.transcribe(tmp_path, language="ko")
            return result.get("text", "").strip()
        finally:
            os.unlink(tmp_path)
    except Exception as exc:
        logger.warning(f"동영상 전사 실패 ({filename}): {exc}")
        return f"[동영상 첨부: {filename}]"


def extract_file_texts(files: list[dict], bot_token: str) -> list[AttachmentResult]:
    """
    Slack files 목록에서 이미지를 제외한 파일의 텍스트를 추출하여 반환한다.

    지원 형식:
      - 텍스트 계열: txt, trf, log, csv, tsv, yaml, json, xml 등
      - 스프레드시트: xlsx, xls (openpyxl 필요)
      - 워드: docx, doc (python-docx 필요)
      - PDF: pdf (pypdf 필요)
      - 동영상: mov, mp4 등 (openai-whisper 선택 설치 시 전사, 미설치 시 파일명만 기록)
      - 미지원: 파일명만 기록

    반환: AttachmentResult 목록. 없으면 빈 리스트.
    """
    results: list[AttachmentResult] = []

    for f in files:
        mime = f.get("mimetype", "")
        filename = f.get("name", "첨부파일")
        file_id = f.get("id", "")
        url = f.get("url_private_download") or f.get("url_private")
        ext = _get_extension(filename)

        if not url:
            continue

        # 이미지는 image_processor.py에서 처리
        if mime.startswith("image/"):
            continue

        extracted = ""
        file_type = "other"

        # ---- MIME 우선 분류, 실패 시 확장자 fallback ----

        # 텍스트 계열
        if any(mime.startswith(p) for p in _TEXT_MIME_PREFIXES) or ext in _TEXT_EXTENSIONS:
            file_type = "text"
            data = _download_bytes(url, bot_token)
            if data:
                extracted = _extract_text_plain(data)

        # xlsx
        elif mime in _XLSX_MIMES or ext in _XLSX_EXTENSIONS:
            file_type = "xlsx"
            data = _download_bytes(url, bot_token)
            if data:
                extracted = _extract_xlsx(data)

        # docx / doc
        elif mime in _DOCX_MIMES or ext in _DOCX_EXTENSIONS:
            file_type = "docx"
            data = _download_bytes(url, bot_token)
            if data:
                extracted = _extract_docx(data)

        # pdf
        elif mime in _PDF_MIMES or ext in _PDF_EXTENSIONS:
            file_type = "pdf"
            data = _download_bytes(url, bot_token)
            if data:
                extracted = _extract_pdf(data)

        # 동영상
        elif mime in _VIDEO_MIMES or ext in _VIDEO_EXTENSIONS:
            file_type = "video"
            data = _download_bytes(url, bot_token)
            if data:
                extracted = _extract_video(data, filename)

        # 미지원 형식 — 파일명만 기록
        else:
            logger.debug(f"미지원 파일 형식 — 이름만 기록: mime={mime!r} name={filename!r}")
            results.append(AttachmentResult(
                slack_file_id=file_id,
                file_name=filename,
                mime_type=mime,
                file_type="other",
                analysis_text=f"[파일: {filename}]",
            ))
            continue

        extracted = extracted[:_MAX_TEXT_CHARS].strip()
        label = f"[파일: {filename}]"
        analysis = f"{label}\n{extracted}" if extracted else label
        results.append(AttachmentResult(
            slack_file_id=file_id,
            file_name=filename,
            mime_type=mime,
            file_type=file_type,
            analysis_text=analysis,
        ))

    return results
